from fastapi import APIRouter
from pydantic import BaseModel
from pathlib import Path
import subprocess
import os

router = APIRouter()

SAVE_DIR = Path.home() / "JarvisFiles"
SAVE_DIR.mkdir(exist_ok=True)

class WordRequest(BaseModel):
    filename: str
    content: str
    folder: str = ""

class ExcelRequest(BaseModel):
    filename: str
    data: list
    headers: list = []
    folder: str = ""

class AppRequest(BaseModel):
    app: str

@router.post("/word/create")
async def create_word(req: WordRequest):
    from docx import Document
    
    folder = SAVE_DIR / req.folder if req.folder else SAVE_DIR
    folder.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    doc.add_heading(req.filename, 0)
    
    for paragraph in req.content.split("\n"):
        if paragraph.strip():
            doc.add_paragraph(paragraph)
    
    filepath = folder / f"{req.filename}.docx"
    doc.save(filepath)
    
    return {"status": "creado", "path": str(filepath)}

@router.post("/excel/create")
async def create_excel(req: ExcelRequest):
    from openpyxl import Workbook
    
    folder = SAVE_DIR / req.folder if req.folder else SAVE_DIR
    folder.mkdir(parents=True, exist_ok=True)
    
    wb = Workbook()
    ws = wb.active
    ws.title = req.filename
    
    if req.headers:
        ws.append(req.headers)
    
    for row in req.data:
        ws.append(row)
    
    filepath = folder / f"{req.filename}.xlsx"
    wb.save(filepath)
    
    return {"status": "creado", "path": str(filepath)}

@router.post("/open")
async def open_file(req: AppRequest):
    app = req.app.lower()
    
    apps = {
        "word": "winword",
        "excel": "excel",
        "notepad": "notepad",
        "chrome": "chrome",
        "spotify": "spotify",
        "calculator": "calc",
    }
    
    cmd = apps.get(app, app)
    
    try:
        subprocess.Popen(cmd, shell=True)
        return {"status": "abierto", "app": app}
    except Exception as e:
        return {"status": "error", "detail": str(e)}

@router.get("/files")
async def list_files():
    files = []
    for f in SAVE_DIR.rglob("*"):
        if f.is_file():
            files.append({
                "name": f.name,
                "path": str(f),
                "size": f.stat().st_size
            })
    return files